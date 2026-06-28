"""Integration tests for candidate management pipeline API."""

from io import BytesIO

import pytest
from docx import Document
from httpx import AsyncClient

REGISTER_PAYLOAD = {
    "organization_name": "Candidate Test Agency",
    "email": "admin@candidatetest.com",
    "password": "SecurePass123",
    "first_name": "Recruiter",
    "last_name": "Admin",
}

CLIENT_PAYLOAD = {
    "name": "HireCorp",
    "status": "active",
    "contacts": [
        {
            "first_name": "HR",
            "last_name": "Manager",
            "email": "hr@hirecorp.com",
            "contact_type": "hiring_manager",
        }
    ],
}

JOB_PAYLOAD = {
    "title": "Senior Python Developer",
    "employment_type": "full_time",
    "priority": "high",
    "status": "open",
    "required_skills": ["Python", "FastAPI"],
    "preferred_skills": ["Docker"],
    "description": "Build backend APIs with Python and FastAPI.",
}

CANDIDATE_PAYLOAD = {
    "first_name": "Jane",
    "last_name": "Doe",
    "email": "jane.doe@example.com",
    "phone": "+1-555-0199",
    "current_company": "TechCorp",
}


def _make_resume_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Jane Doe — Senior Python Developer")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def _setup_full_context(client: AsyncClient) -> tuple[str, str, str]:
    reg = await client.post("/api/v1/auth/register", json=REGISTER_PAYLOAD)
    token = reg.json()["tokens"]["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    client_resp = await client.post("/api/v1/clients", json=CLIENT_PAYLOAD, headers=headers)
    client_id = client_resp.json()["id"]
    job_resp = await client.post(
        "/api/v1/jobs",
        json={**JOB_PAYLOAD, "client_id": client_id},
        headers=headers,
    )
    job_id = job_resp.json()["id"]
    candidate_resp = await client.post(
        "/api/v1/candidates",
        json={**CANDIDATE_PAYLOAD, "job_requirement_id": job_id},
        headers=headers,
    )
    candidate_id = candidate_resp.json()["id"]
    return token, candidate_id, job_id


@pytest.mark.asyncio
class TestCandidateAPI:
    async def test_create_candidate_with_application(self, client: AsyncClient) -> None:
        token, candidate_id, job_id = await _setup_full_context(client)
        headers = {"Authorization": f"Bearer {token}"}
        response = await client.get(f"/api/v1/candidates/{candidate_id}", headers=headers)
        assert response.status_code == 200
        assert response.json()["full_name"] == "Jane Doe"

        apps = await client.get(
            "/api/v1/applications",
            params={"candidate_id": candidate_id},
            headers=headers,
        )
        assert apps.status_code == 200
        assert apps.json()["total"] == 1
        assert apps.json()["items"][0]["job_requirement_id"] == job_id

    async def test_full_pipeline(self, client: AsyncClient) -> None:
        token, candidate_id, job_id = await _setup_full_context(client)
        headers = {"Authorization": f"Bearer {token}"}

        apps = await client.get(
            "/api/v1/applications",
            params={"candidate_id": candidate_id},
            headers=headers,
        )
        app_id = apps.json()["items"][0]["id"]

        resume_bytes = _make_resume_docx()
        upload = await client.post(
            f"/api/v1/applications/{app_id}/resume",
            headers=headers,
            files={"file": ("resume.docx", resume_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert upload.status_code == 200
        detail = upload.json()
        assert detail["parsed_resume"]["status"] == "completed"
        assert detail["pipeline_stage"] == "screening"

        context = await client.get(
            f"/api/v1/applications/{app_id}/resume-build-context",
            headers=headers,
        )
        assert context.status_code == 200
        assert "candidate_total_experience_years" in context.json()
        assert "gap_strategy_options" in context.json()

        build = await client.post(
            f"/api/v1/applications/{app_id}/build-resume",
            json={"template_id": context.json()["default_template_id"]},
            headers=headers,
        )
        assert build.status_code == 200
        version_id = build.json()["id"]
        assert build.json()["status"] == "pending_review"
        assert build.json()["template_id"] is not None
        assert build.json()["template_name"] is not None

        review = await client.post(
            f"/api/v1/applications/{app_id}/resume-versions/{version_id}/review",
            json={"decision": "accept", "notes": "Looks good"},
            headers=headers,
        )
        assert review.status_code == 200
        assert review.json()["status"] == "approved"

        score = await client.post(
            f"/api/v1/applications/{app_id}/score-resume",
            headers=headers,
        )
        assert score.status_code == 200
        assert float(score.json()["overall_score"]) > 0

        export = await client.get(
            f"/api/v1/applications/{app_id}/export/docx",
            headers=headers,
        )
        assert export.status_code == 200
        assert len(export.content) > 0

        schedule = await client.post(
            f"/api/v1/applications/{app_id}/schedule-interview",
            json={"duration_minutes": 30, "difficulty": "medium"},
            headers=headers,
        )
        assert schedule.status_code == 200
        assert schedule.json()["status"] == "scheduled"

        start = await client.post(
            f"/api/v1/applications/{app_id}/interview/start",
            headers=headers,
        )
        assert start.status_code == 200
        assert start.json()["question_number"] == 1

        answer = await client.post(
            f"/api/v1/applications/{app_id}/interview/answer",
            json={"answer": "I have 5 years of Python experience building APIs."},
            headers=headers,
        )
        assert answer.status_code == 200

        complete = await client.post(
            f"/api/v1/applications/{app_id}/interview/complete",
            headers=headers,
        )
        assert complete.status_code == 200
        assert complete.json()["status"] == "completed"
        assert complete.json()["summary"]
        assert complete.json()["recommendation"] == "hire"

        pipeline = await client.get(
            "/api/v1/applications",
            params={"candidate_id": candidate_id},
            headers=headers,
        )
        summary = pipeline.json()["items"][0]
        assert summary["has_parsed_resume"] is True
        assert summary["resume_version_count"] == 1
        assert summary["latest_resume_version_id"] == version_id
        assert float(summary["latest_score"]) > 0
        assert summary["interview_status"] == "completed"
        assert float(summary["interview_overall_score"]) > 0

    async def test_list_candidates(self, client: AsyncClient) -> None:
        token, _, _ = await _setup_full_context(client)
        headers = {"Authorization": f"Bearer {token}"}
        response = await client.get("/api/v1/candidates", headers=headers)
        assert response.status_code == 200
        assert response.json()["total"] >= 1
