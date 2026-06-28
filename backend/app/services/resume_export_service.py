"""Resume export to PDF and DOCX."""

import json
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.core.exceptions import ValidationError
from app.core.logging import get_logger

logger = get_logger(__name__)

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates" / "resume"


class ResumeExportService:
    """Render resume JSON to HTML, PDF (Playwright), and DOCX."""

    def __init__(self) -> None:
        self._env = Environment(
            loader=FileSystemLoader(str(_TEMPLATES_DIR)),
            autoescape=select_autoescape(["html", "xml"]),
        )
        self._string_env = Environment(autoescape=select_autoescape(["html", "xml"]))

    def render_html_from_template(
        self,
        content_json: dict[str, Any],
        html_template: str,
        css_styles: str = "",
    ) -> str:
        template = self._string_env.from_string(html_template)
        body = template.render(resume=content_json)
        stripped = html_template.strip().lower()
        if stripped.startswith("<!doctype") or stripped.startswith("<html"):
            return body
        css_block = f"<style>{css_styles}</style>" if css_styles else ""
        return (
            f"<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"UTF-8\"/>{css_block}</head>"
            f"<body>{body}</body></html>"
        )

    def render_html(
        self,
        content_json: dict[str, Any],
        template_name: str = "default.html",
        html_template: str | None = None,
        css_styles: str | None = None,
    ) -> str:
        if html_template is not None:
            return self.render_html_from_template(content_json, html_template, css_styles or "")
        template = self._env.get_template(template_name)
        return template.render(resume=content_json)

    def export_docx(self, content_json: dict[str, Any]) -> bytes:
        """Generate DOCX from resume JSON."""
        doc = Document()
        basic = content_json.get("basic_details", {})
        name = basic.get("full_name", "Resume")
        doc.add_heading(name, 0)

        if basic.get("email"):
            doc.add_paragraph(f"Email: {basic['email']}")
        if basic.get("phone"):
            doc.add_paragraph(f"Phone: {basic['phone']}")
        if basic.get("location"):
            doc.add_paragraph(f"Location: {basic['location']}")

        if content_json.get("summary"):
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(content_json["summary"])

        if content_json.get("experience"):
            doc.add_heading("Experience", level=1)
            for exp in content_json["experience"]:
                title_line = f"{exp.get('title', '')} — {exp.get('company', '')}"
                doc.add_paragraph(title_line, style="List Bullet")
                for bullet in exp.get("bullets", []):
                    p = doc.add_paragraph(bullet, style="List Bullet 2")
                    p.runs[0].font.size = Pt(10)

        if content_json.get("skills"):
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(content_json["skills"]))

        if content_json.get("education"):
            doc.add_heading("Education", level=1)
            for edu in content_json["education"]:
                doc.add_paragraph(
                    f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('graduation_year', '')})"
                )

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def export_html(
        self,
        content_json: dict[str, Any],
        html_template: str | None = None,
        css_styles: str | None = None,
    ) -> bytes:
        return self.render_html(content_json, html_template=html_template, css_styles=css_styles).encode("utf-8")

    def export_pdf(
        self,
        content_json: dict[str, Any],
        html_template: str | None = None,
        css_styles: str | None = None,
    ) -> bytes:
        """Generate PDF from HTML (xhtml2pdf; Playwright fallback if installed)."""
        html = self.render_html(content_json, html_template=html_template, css_styles=css_styles)
        buffer = BytesIO()

        try:
            from xhtml2pdf import pisa

            result = pisa.CreatePDF(html, dest=buffer, encoding="utf-8")
            if result.err:
                raise ValidationError(f"PDF generation failed: {result.err}")
            pdf_bytes = buffer.getvalue()
            if pdf_bytes:
                return pdf_bytes
        except ImportError:
            logger.warning("xhtml2pdf_not_installed")
        except ValidationError:
            raise
        except Exception as exc:
            logger.warning("xhtml2pdf_failed", error=str(exc))

        try:
            from playwright.sync_api import sync_playwright
        except ImportError as exc:
            raise ValidationError(
                "PDF export failed. Install xhtml2pdf or playwright in the backend environment."
            ) from exc

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.set_content(html, wait_until="networkidle")
                pdf_bytes = page.pdf(format="A4", print_background=True)
                browser.close()
                return pdf_bytes
        except Exception as exc:
            logger.error("pdf_export_failed", error=str(exc))
            raise ValidationError(f"PDF generation failed: {exc}") from exc

    def export_json_bytes(self, content_json: dict[str, Any]) -> bytes:
        return json.dumps(content_json, indent=2).encode("utf-8")
